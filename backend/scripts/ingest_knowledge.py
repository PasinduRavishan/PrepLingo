"""Ingest knowledge documents of multiple file formats into ChromaDB.

Supported extensions:
- .md
- .txt
- .pdf
- .docx
- .html/.htm

How to run:
  cd PrepLingo/backend
  python scripts/ingest_knowledge.py

Optional:
  python scripts/ingest_knowledge.py --reset

Notes:
1. Keep curated markdown in ./knowledge_base
2. Add raw docs/articles/pdfs in ./knowledge_raw (optional)
3. Re-running ingestion replaces chunks for touched sources, preventing duplicates
"""

import argparse
import hashlib
import os
import re
import sys
import time
from collections import Counter, defaultdict
from pathlib import Path

import chromadb
import fitz
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

# Add parent directory to path so we can import from app/
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

load_dotenv()

BACKEND_DIR = Path(__file__).resolve().parent.parent
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
raw_vector_store_path = os.getenv("VECTOR_STORE_PATH", "./vector_store_data")
VECTOR_STORE_PATH = str((BACKEND_DIR / raw_vector_store_path).resolve())
KNOWLEDGE_BASE_PATH = BACKEND_DIR / "knowledge_base"
KNOWLEDGE_RAW_PATH = BACKEND_DIR / "knowledge_raw"
SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".html", ".htm"}

# Filename stems to always skip in knowledge_raw (admin/meta files, not interview content).
_SKIP_STEMS: frozenset[str] = frozenset({
    "readme",
    "collector_readme",
    "discovery_report",
    "seed_sources",
})

# Filename prefixes to skip in knowledge_raw.
# "seed_"  → duplicate seed URLs saved before the crawler ran
# "index_" → site index/listing pages (nav links only, no interview content)
_SKIP_PREFIXES: tuple[str, ...] = ("seed_", "index_")


def determine_interview_type(source_path: str) -> str:
    """Map file path to interview type metadata."""
    path_lower = source_path.lower().replace("\\", "/")
    if "/technical/" in path_lower:
        return "technical"
    if "/system_design/" in path_lower:
        return "system_design"
    if "/behavioral/" in path_lower:
        return "behavioral"
    if "/resume_interview/" in path_lower:
        return "resume"
    return "general"


def determine_topic(source_path: str) -> str:
    """Infer topic from the segment after interview type, fallback to file stem."""
    clean = source_path.replace("\\", "/")
    parts = [p for p in clean.split("/") if p]
    known_types = {"technical", "system_design", "behavioral", "resume_interview"}
    for idx, part in enumerate(parts):
        if part in known_types:
            if idx + 1 < len(parts) - 1:
                return parts[idx + 1]
            if idx + 1 < len(parts):
                return Path(parts[idx + 1]).stem
    return Path(source_path).stem


def _should_skip_raw_file(path: Path) -> bool:
    """
    Return True for files in knowledge_raw that are noise, not interview content.

    Skipped categories:
      - Admin/meta files: README.md, collector_README.md, discovery_report.md
      - Seed duplicates:  seed_*.html  (collector saved seed URLs before crawling;
                          the crawled versions are also present — ingesting both
                          doubles chunk count with identical content)
      - Index pages:      index_*.html  (site listing/nav pages, no real content)
    """
    stem = path.stem.lower()
    return stem in _SKIP_STEMS or any(stem.startswith(p) for p in _SKIP_PREFIXES)


def discover_files(root: Path, is_raw: bool = False) -> list[Path]:
    """Find supported knowledge files under root recursively.

    When is_raw=True (i.e. scanning knowledge_raw/), apply filename filters
    to exclude meta files, seed duplicates, and index pages.
    """
    if not root.exists():
        return []
    files: list[Path] = []
    for path in root.rglob("*"):
        if not (path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS):
            continue
        if is_raw and _should_skip_raw_file(path):
            print(f"   ⊘ Skipping noise file: {path.name}")
            continue
        files.append(path)
    return sorted(files)


def read_file_text(path: Path) -> str:
    """Load text content based on file extension."""
    ext = path.suffix.lower()

    if ext in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    if ext == ".pdf":
        doc = fitz.open(path)
        try:
            return "\n".join(page.get_text("text") for page in doc)
        finally:
            doc.close()

    if ext == ".docx":
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

    if ext in {".html", ".htm"}:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n")

    raise ValueError(f"Unsupported file extension: {ext}")


def normalize_text(text: str) -> str:
    """Collapse noisy whitespace while keeping paragraphs."""
    lines = [line.strip() for line in text.splitlines()]
    cleaned = [line for line in lines if line]
    return "\n".join(cleaned).strip()


def relative_source(path: Path, roots: list[Path]) -> str:
    """Create stable, readable source path with root prefix."""
    for root in roots:
        try:
            rel = path.relative_to(root)
            return f"{root.name}/{rel.as_posix()}"
        except ValueError:
            continue
    return path.as_posix()


def build_documents(files: list[Path], roots: list[Path]) -> list[Document]:
    """Create LangChain Document objects with enriched metadata.

    Deduplicates by content SHA-256 so that near-identical HTML pages
    (e.g. two cached versions of the same AWS article) only produce one
    document each. The first file encountered wins; later duplicates are
    logged and skipped.
    """
    docs: list[Document] = []
    seen_hashes: dict[str, str] = {}  # hash → first source path

    for file_path in files:
        try:
            raw = read_file_text(file_path)
            text = normalize_text(raw)
            if not text:
                print(f"   ⚠ Skipping empty file: {file_path}")
                continue

            content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
            if content_hash in seen_hashes:
                print(f"   ⊘ Skipping duplicate of {seen_hashes[content_hash]}: {file_path.name}")
                continue
            seen_hashes[content_hash] = str(file_path)

            source = relative_source(file_path, roots)
            metadata = {
                "source": source,
                "source_type": file_path.suffix.lower().lstrip("."),
                "interview_type": determine_interview_type(source),
                "topic": determine_topic(source),
                "title": file_path.stem,
                "content_hash": content_hash,
            }
            docs.append(Document(page_content=text, metadata=metadata))
        except Exception as exc:
            print(f"   ⚠ Failed to load {file_path}: {exc}")
    return docs


def reset_knowledge_collection(persist_path: str):
    """Delete only the knowledge collection while preserving other collections."""
    client = chromadb.PersistentClient(path=persist_path)
    try:
        client.delete_collection("knowledge")
        print("   ✅ Reset existing 'knowledge' collection")
    except Exception:
        print("   ℹ No existing 'knowledge' collection to reset")


def chunk_documents(documents: list[Document]) -> list[Document]:
    """Split each source document into retrievable chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)

    # Attach per-source chunk index and chunk hash for traceability.
    per_source_index: defaultdict[str, int] = defaultdict(int)
    for chunk in chunks:
        source = chunk.metadata.get("source", "unknown")
        per_source_index[source] += 1
        chunk.metadata["chunk_index"] = per_source_index[source]
        chunk.metadata["chunk_hash"] = hashlib.sha256(
            chunk.page_content.encode("utf-8")
        ).hexdigest()
    return chunks


def build_chunk_ids(chunks: list[Document]) -> list[str]:
    """Generate stable IDs so upserts are deterministic."""
    ids: list[str] = []
    for chunk in chunks:
        source = chunk.metadata.get("source", "unknown")
        chunk_index = chunk.metadata.get("chunk_index", 0)
        chunk_hash = chunk.metadata.get("chunk_hash", "")
        raw = f"{source}:{chunk_index}:{chunk_hash}"
        ids.append(hashlib.sha1(raw.encode("utf-8")).hexdigest())
    return ids


def parse_retry_delay_seconds(error_text: str, default_delay: float = 25.0) -> float:
    """Extract retry delay from provider error text when available."""
    patterns = [
        r"retry in\s+([0-9]+(?:\.[0-9]+)?)s",
        r"retryDelay['\"]?\s*:\s*['\"]?([0-9]+)s",
    ]
    text = error_text or ""
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            try:
                return max(float(match.group(1)), 1.0)
            except Exception:
                pass
    return default_delay


def is_quota_or_rate_limit_error(exc: Exception) -> bool:
    """Detect provider quota/rate-limit error signatures."""
    text = str(exc).lower()
    signals = [
        "resource_exhausted",
        "quota exceeded",
        "429",
        "rate limit",
        "too many requests",
    ]
    return any(signal in text for signal in signals)


def add_chunks_with_retry(
    vector_store: Chroma,
    chunks: list[Document],
    chunk_ids: list[str],
    batch_size: int,
    max_retries: int,
    sleep_between_batches: float,
) -> None:
    """Add chunks in batches with resilient retry/backoff on 429 errors."""
    total = len(chunks)
    if total == 0:
        return

    for start in range(0, total, batch_size):
        end = min(start + batch_size, total)
        batch_docs = chunks[start:end]
        batch_ids = chunk_ids[start:end]
        attempt = 0

        while True:
            attempt += 1
            try:
                vector_store.add_documents(documents=batch_docs, ids=batch_ids)
                print(f"   ✅ Batch {start + 1}-{end}/{total} stored")
                if sleep_between_batches > 0:
                    time.sleep(sleep_between_batches)
                break
            except Exception as exc:
                if not is_quota_or_rate_limit_error(exc):
                    raise

                if attempt > max_retries:
                    raise RuntimeError(
                        f"Exceeded retry limit while embedding batch {start + 1}-{end}"
                    ) from exc

                wait_seconds = parse_retry_delay_seconds(str(exc)) + 1.5
                print(
                    f"   ⚠ Rate limit on batch {start + 1}-{end}. "
                    f"Retry {attempt}/{max_retries} in {wait_seconds:.1f}s"
                )
                time.sleep(wait_seconds)


def main(
    reset: bool = False,
    batch_size: int = 32,
    max_retries: int = 6,
    sleep_between_batches: float = 0.8,
):
    print("=" * 64)
    print("PrepLingo Knowledge Ingestion (Multi-Format)")
    print("=" * 64)

    # GOOGLE_API_KEY is no longer required for embeddings (switched to local HuggingFace).
    # Keep the check as a warning only in case other parts of the system still use it.

    roots = [KNOWLEDGE_BASE_PATH]
    if KNOWLEDGE_RAW_PATH.exists():
        roots.append(KNOWLEDGE_RAW_PATH)

    print("\nScanning knowledge roots:")
    for root in roots:
        print(f" - {root}")

    all_files: list[Path] = []
    for root in roots:
        is_raw = root == KNOWLEDGE_RAW_PATH
        all_files.extend(discover_files(root, is_raw=is_raw))

    print(f"\nFound {len(all_files)} supported files")
    if not all_files:
        print("No files found. Add docs under knowledge_base/ or knowledge_raw/")
        sys.exit(1)

    print("\nLoading and normalizing documents...")
    documents = build_documents(all_files, roots)
    print(f"   ✅ Loaded {len(documents)} non-empty documents")
    if not documents:
        print("No usable documents after parsing. Check source files.")
        sys.exit(1)

    type_counts = Counter(doc.metadata["interview_type"] for doc in documents)
    format_counts = Counter(doc.metadata["source_type"] for doc in documents)
    print("\nDocument breakdown by interview type:")
    for key, count in sorted(type_counts.items()):
        print(f" - {key}: {count}")
    print("\nDocument breakdown by file type:")
    for key, count in sorted(format_counts.items()):
        print(f" - {key}: {count}")

    print("\nChunking documents...")
    chunks = chunk_documents(documents)
    print(f"   ✅ Created {len(chunks)} chunks")

    print("\nInitializing embeddings (local HuggingFace — no API key needed)...")
    print("   Model: BAAI/bge-base-en-v1.5 (768 dims, MTEB 72.3)")
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")

    if reset:
        print("\nReset requested")
        reset_knowledge_collection(VECTOR_STORE_PATH)

    vector_store = Chroma(
        collection_name="knowledge",
        embedding_function=embeddings,
        persist_directory=VECTOR_STORE_PATH,
    )

    # Replace chunks per source to avoid duplicate growth on reruns.
    print("\nReplacing existing chunks for touched sources...")
    touched_sources = sorted({c.metadata.get("source", "") for c in chunks if c.metadata.get("source")})
    for source in touched_sources:
        try:
            vector_store.delete(where={"source": source})
        except Exception:
            pass

    print("Adding new chunks to ChromaDB...")
    chunk_ids = build_chunk_ids(chunks)
    add_chunks_with_retry(
        vector_store=vector_store,
        chunks=chunks,
        chunk_ids=chunk_ids,
        batch_size=max(batch_size, 1),
        max_retries=max(max_retries, 1),
        sleep_between_batches=max(sleep_between_batches, 0.0),
    )
    print(f"   ✅ Stored {len(chunks)} chunks in 'knowledge' collection")

    print("\nSmoke-testing retrieval...")
    retriever = vector_store.as_retriever(
        search_kwargs={"filter": {"interview_type": "technical"}, "k": 2}
    )
    test_results = retriever.invoke("How does database indexing work?")
    print(f"Retrieved {len(test_results)} chunks for technical query")
    for idx, doc in enumerate(test_results, start=1):
        snippet = doc.page_content[:100].replace("\n", " ")
        print(f" [{idx}] {doc.metadata.get('source')} | {snippet}...")

    print("\n" + "=" * 64)
    print("Knowledge ingestion complete")
    print("Now run: uvicorn app.main:app --reload")
    print("=" * 64)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingest multi-format knowledge into ChromaDB")
    parser.add_argument(
        "--reset",
        action="store_true",
        help="Delete existing knowledge collection before ingest",
    )
    parser.add_argument(
        "--batch-size",
        type=int,
        default=32,
        help="Number of chunks to embed/write per request batch",
    )
    parser.add_argument(
        "--max-retries",
        type=int,
        default=6,
        help="Maximum retries per batch when rate-limited",
    )
    parser.add_argument(
        "--sleep-between-batches",
        type=float,
        default=0.8,
        help="Seconds to pause between successful batches",
    )
    args = parser.parse_args()
    main(
        reset=args.reset,
        batch_size=args.batch_size,
        max_retries=args.max_retries,
        sleep_between_batches=args.sleep_between_batches,
    )
